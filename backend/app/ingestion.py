import base64
import concurrent.futures
import hashlib
import json
import logging
import re
import urllib.request
from pathlib import Path

import fitz
import pytesseract
from docx import Document as DocxDocument
from PIL import Image, ImageFilter, ImageOps

from app.config import Settings
from app.models import ChunkRecord, IngestDocument, TextUnit
from app.text_sanitize import sanitize_pg_text

logger = logging.getLogger(__name__)

SUPPORTED_EXTENSIONS = {".txt", ".pdf", ".docx"}


def safe_document_id(path: Path, docs_dir: Path) -> str:
    return path.relative_to(docs_dir).as_posix()


def content_hash(content: str) -> str:
    return hashlib.sha256(content.encode("utf-8")).hexdigest()


def _read_txt(path: Path) -> str:
    for encoding in ("utf-8", "utf-8-sig", "cp1251"):
        try:
            return path.read_text(encoding=encoding)
        except UnicodeDecodeError:
            continue
    return path.read_text(encoding="utf-8", errors="ignore")


def _read_docx(path: Path) -> str:
    doc = DocxDocument(str(path))
    lines: list[str] = ["[PAGE 1]"]
    page = 1
    chars = 0
    for paragraph in doc.paragraphs:
        text = (paragraph.text or "").strip()
        if not text:
            continue
        if chars >= 2200:
            page += 1
            chars = 0
            lines.append(f"[PAGE {page}]")
        lines.append(text)
        chars += len(text)
    return "\n".join(lines)


def _normalize_ocr_output(text: str) -> str:
    table = str.maketrans(
        {
            "A": "А",
            "B": "В",
            "C": "С",
            "E": "Е",
            "H": "Н",
            "K": "К",
            "M": "М",
            "O": "О",
            "P": "Р",
            "T": "Т",
            "X": "Х",
            "a": "а",
            "c": "с",
            "e": "е",
            "o": "о",
            "p": "р",
            "x": "х",
            "y": "у",
        }
    )
    normalized = text.translate(table)
    normalized = re.sub(r"[ \t]+", " ", normalized)
    normalized = re.sub(r"\n{3,}", "\n\n", normalized)
    return normalized.strip()


def _text_quality_score(text: str) -> float:
    cleaned = re.sub(r"\s+", " ", text).strip()
    if not cleaned:
        return 0.0
    alpha = len(re.findall(r"[A-Za-zА-Яа-яЁё]", cleaned))
    cyr = len(re.findall(r"[А-Яа-яЁё]", cleaned))
    long_words = len(re.findall(r"[A-Za-zА-Яа-яЁё]{4,}", cleaned))
    noise = len(re.findall(r"[^A-Za-zА-Яа-яЁё0-9\s.,:;!?()\"'«»\-]", cleaned))
    return (
        min(len(cleaned), 8000) / 8000.0
        + (alpha / max(1, len(cleaned))) * 1.4
        + (cyr / max(1, alpha)) * 1.2
        + min(long_words, 200) / 200.0
        - (noise / max(1, len(cleaned))) * 3.0
    )


def _page_png_bytes(page, zoom: float) -> bytes:
    pix = page.get_pixmap(matrix=fitz.Matrix(zoom, zoom), alpha=False)
    return pix.tobytes("png")


def _ocr_page(pdf_doc, page_index: int, lang: str) -> str:
    page = pdf_doc.load_page(page_index)
    pix = page.get_pixmap(matrix=fitz.Matrix(3, 3), alpha=False)
    base = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
    gray = ImageOps.grayscale(base)
    contrast = ImageOps.autocontrast(gray)
    denoised = contrast.filter(ImageFilter.MedianFilter(size=3))
    binary = denoised.point(lambda p: 255 if p > 175 else 0)
    variants = [
        (denoised, "--oem 1 --psm 3 -c preserve_interword_spaces=1"),
        (denoised, "--oem 1 --psm 4 -c preserve_interword_spaces=1"),
        (binary, "--oem 1 --psm 6 -c preserve_interword_spaces=1"),
    ]
    best = ""
    best_score = -10**9
    for image, config in variants:
        candidate = _normalize_ocr_output(pytesseract.image_to_string(image, lang=lang, config=config))
        score = _text_quality_score(candidate)
        if score > best_score:
            best = candidate
            best_score = score
    return best


def _call_vl(settings: Settings, page_png: bytes) -> str:
    if not settings.vl_api_url:
        return ""
    headers = {"Content-Type": "application/json"}
    if settings.vl_api_key:
        headers["Authorization"] = f"Bearer {settings.vl_api_key}"
    image_b64 = base64.b64encode(page_png).decode("ascii")
    payload = {
        "model": settings.vl_model,
        "messages": [
            {
                "role": "user",
                "content": [
                    {
                        "type": "text",
                        "text": (
                            "Extract all readable text from this document page. "
                            "Return plain text only, preserve Russian text when present."
                        ),
                    },
                    {"type": "image_url", "image_url": {"url": f"data:image/png;base64,{image_b64}"}},
                ],
            }
        ],
        "temperature": 0,
    }
    request = urllib.request.Request(
        settings.vl_api_url,
        data=json.dumps(payload).encode("utf-8"),
        method="POST",
        headers=headers,
    )
    try:
        with urllib.request.urlopen(request, timeout=settings.vl_timeout_sec) as response:
            body = json.loads(response.read().decode("utf-8"))
        return str(body["choices"][0]["message"]["content"]).strip()
    except Exception as exc:  # noqa: BLE001
        logger.warning("VL extraction failed: %s", exc)
        return ""


def _page_cache_path(settings: Settings, pdf_path: Path, page_num: int) -> Path:
    stat = pdf_path.stat()
    key = hashlib.sha256(f"{pdf_path.resolve()}:{stat.st_mtime}:{stat.st_size}".encode("utf-8")).hexdigest()
    return settings.page_cache_dir / key / f"page_{page_num:05d}.json"


def _read_pdf_page(settings: Settings, pdf_path: Path, page_index: int) -> tuple[int, str]:
    page_num = page_index + 1
    cache_path = _page_cache_path(settings, pdf_path, page_num)
    if cache_path.is_file():
        try:
            cached = json.loads(cache_path.read_text(encoding="utf-8"))
            return page_num, str(cached.get("text", ""))
        except (OSError, json.JSONDecodeError):
            pass

    with fitz.open(str(pdf_path)) as pdf_doc:
        page = pdf_doc.load_page(page_index)
        pdf_text = _normalize_ocr_output(page.get_text() or "")
        pdf_score = _text_quality_score(pdf_text)

        if settings.prepare_page_engine == "pdf_text":
            chosen = pdf_text
        elif settings.prepare_page_engine == "tesseract":
            ocr_text = _ocr_page(pdf_doc, page_index, settings.tesseract_lang)
            chosen = ocr_text if _text_quality_score(ocr_text) >= pdf_score else pdf_text
        elif settings.prepare_page_engine == "vl":
            chosen = _call_vl(settings, _page_png_bytes(page, settings.vl_zoom))
        else:
            if pdf_score >= settings.vl_min_quality and len(pdf_text) >= 80:
                chosen = pdf_text
            else:
                ocr_text = _ocr_page(pdf_doc, page_index, settings.tesseract_lang)
                ocr_score = _text_quality_score(ocr_text)
                chosen = ocr_text if ocr_score >= pdf_score else pdf_text
                if _text_quality_score(chosen) < settings.vl_min_quality and settings.vl_api_url:
                    vl_text = _call_vl(settings, _page_png_bytes(page, settings.vl_zoom))
                    if _text_quality_score(vl_text) > _text_quality_score(chosen):
                        chosen = vl_text

    cache_path.parent.mkdir(parents=True, exist_ok=True)
    cache_path.write_text(json.dumps({"text": chosen}, ensure_ascii=False), encoding="utf-8")
    return page_num, chosen


def _read_pdf(settings: Settings, path: Path) -> str:
    with fitz.open(str(path)) as pdf_doc:
        total_pages = pdf_doc.page_count
    pages: list[tuple[int, str]] = []
    with concurrent.futures.ThreadPoolExecutor(max_workers=settings.prepare_workers) as executor:
        futures = [executor.submit(_read_pdf_page, settings, path, page_index) for page_index in range(total_pages)]
        for future in concurrent.futures.as_completed(futures):
            pages.append(future.result())
    pages.sort(key=lambda item: item[0])
    return "\n\n".join(f"[PAGE {page_num}]\n{text}" for page_num, text in pages if text.strip())


def read_document_content(settings: Settings, path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix == ".txt":
        content = _read_txt(path)
    elif suffix == ".docx":
        content = _read_docx(path)
    elif suffix == ".pdf":
        content = _read_pdf(settings, path)
    else:
        content = ""
    return sanitize_pg_text(content) or ""


def _extract_units(content: str) -> list[TextUnit]:
    units: list[TextUnit] = []
    current_page: str | None = None
    buffer: list[str] = []

    def flush() -> None:
        text = "\n".join(buffer).strip()
        if text:
            units.append(TextUnit(text=text, page=current_page))
        buffer.clear()

    for line in content.splitlines():
        match = re.match(r"^\s*\[PAGE\s+(\d+)\]\s*$", line, flags=re.IGNORECASE)
        if match:
            flush()
            current_page = match.group(1)
            continue
        buffer.append(line)
    flush()
    return units


def read_document(settings: Settings, path: Path) -> IngestDocument | None:
    if path.suffix.lower() not in SUPPORTED_EXTENSIONS:
        return None
    content = read_document_content(settings, path)
    if not content.strip():
        return None
    units = _extract_units(content)
    if not units:
        units = [TextUnit(text=content, page=None)]
    return IngestDocument(
        document_id=safe_document_id(path, settings.docs_dir),
        file_name=path.name,
        content=content,
        units=units,
    )


def read_documents(settings: Settings) -> list[IngestDocument]:
    if not settings.docs_dir.exists():
        return []
    docs: list[IngestDocument] = []
    for path in sorted(settings.docs_dir.rglob("*")):
        if path.is_file() and path.suffix.lower() in SUPPORTED_EXTENSIONS:
            doc = read_document(settings, path)
            if doc is not None:
                docs.append(doc)
    return docs


def _prefer_chunk_end(text: str, start: int, raw_end: int, chunk_size: int) -> int:
    if raw_end >= len(text):
        return raw_end
    window = text[start:raw_end]
    min_keep = max(80, chunk_size // 3)
    for sep in ("\n\n", "\n", ". "):
        split_at = window.rfind(sep)
        if split_at >= min_keep:
            return start + split_at + len(sep)
    return raw_end


def chunk_text(text: str, chunk_size: int, overlap: int) -> list[str]:
    if not text or chunk_size <= 0:
        return []
    overlap = max(0, min(overlap, chunk_size - 1))
    chunks: list[str] = []
    start = 0
    while start < len(text):
        end = _prefer_chunk_end(text, start, min(len(text), start + chunk_size), chunk_size)
        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)
        if end >= len(text):
            break
        start = max(start + 1, end - overlap)
    return chunks


def chunk_document(settings: Settings, document: IngestDocument) -> list[ChunkRecord]:
    chunks: list[ChunkRecord] = []
    for unit in document.units:
        paragraphs = [part.strip() for part in re.split(r"\n\s*\n+", unit.text) if part.strip()]
        pieces: list[str] = []
        buffer = ""
        for paragraph in paragraphs or [unit.text]:
            candidate = f"{buffer}\n\n{paragraph}".strip() if buffer else paragraph
            if len(candidate) <= settings.chunk_size:
                buffer = candidate
            else:
                if buffer:
                    pieces.append(buffer)
                pieces.extend(chunk_text(paragraph, settings.chunk_size, 0))
                buffer = ""
        if buffer:
            pieces.append(buffer)
        for piece in pieces:
            for chunk in chunk_text(piece, settings.chunk_size, settings.chunk_overlap):
                chunks.append(ChunkRecord(text=sanitize_pg_text(chunk) or "", page=unit.page))
    if settings.chunk_min_merge_chars > 0 and len(chunks) > 1:
        chunks = _merge_tiny_chunks(chunks, settings.chunk_min_merge_chars)
    return chunks


def _merge_tiny_chunks(chunks: list[ChunkRecord], min_chars: int) -> list[ChunkRecord]:
    merged: list[ChunkRecord] = []
    acc = chunks[0]
    for nxt in chunks[1:]:
        if len(acc.text) < min_chars and acc.page == nxt.page:
            acc = ChunkRecord(
                text=sanitize_pg_text(f"{acc.text}\n\n{nxt.text}".strip()) or "",
                page=acc.page,
            )
        else:
            merged.append(acc)
            acc = nxt
    merged.append(acc)
    return merged
